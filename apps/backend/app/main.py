import io
import json
import logging
import os
import threading
import time
from collections import defaultdict, deque
from pathlib import Path
from typing import Any, AsyncIterator
from urllib.parse import quote
from uuid import uuid4

import requests
from bs4 import BeautifulSoup
from pypdf import PdfReader
from docx import Document as DocxDocument
from starlette.requests import Request
from starlette.responses import JSONResponse, PlainTextResponse, Response, StreamingResponse
from haystack import Document
from haystack.components.embedders import (
    SentenceTransformersDocumentEmbedder,
    SentenceTransformersTextEmbedder,
)
from haystack.components.retrievers.in_memory import (
    InMemoryBM25Retriever,
    InMemoryEmbeddingRetriever,
)
from haystack.document_stores.in_memory import InMemoryDocumentStore
from haystack_integrations.document_stores.qdrant import QdrantDocumentStore
from haystack_integrations.components.retrievers.qdrant import QdrantEmbeddingRetriever
from prometheus_client import CONTENT_TYPE_LATEST, Counter, Histogram, generate_latest
from pythonjsonlogger import jsonlogger
from ray import serve

from app.vllm_client import VllmStreamingGenerator


def configure_logging() -> None:
    handler = logging.StreamHandler()
    formatter = jsonlogger.JsonFormatter()
    handler.setFormatter(formatter)
    root = logging.getLogger()
    root.setLevel(logging.INFO)
    root.handlers = [handler]


def env_flag(name: str, default: str = "false") -> bool:
    return os.getenv(name, default).lower() in {"1", "true", "yes", "on"}


def sse(event: str, data: dict[str, Any]) -> bytes:
    return f"event: {event}\ndata: {json.dumps(data, ensure_ascii=False)}\n\n".encode()


BENCH_REQUIREMENTS = "httpx==0.27.2\n"

BENCH_SCRIPT = """#!/usr/bin/env python3
\"\"\"
Lightweight streaming benchmark for /query/stream.

Measures TTFT, tokens/sec (approx), and total latency.
Portable across Akamai LKE, AWS EKS, and GCP GKE.
\"\"\"

from __future__ import annotations

import argparse
import asyncio
import json
import statistics
import time
from pathlib import Path
from typing import Any, Optional

import httpx


def percentile(values: list[float], pct: float) -> float:
    if not values:
        return 0.0
    ordered = sorted(values)
    index = max(0, int(len(ordered) * pct) - 1)
    return ordered[index]


async def run_request(
    client: httpx.AsyncClient,
    url: str,
    prompt: str,
    request_id: int,
) -> dict[str, Any]:
    payload = {"query": prompt}
    start = time.perf_counter()
    ttft: Optional[float] = None
    token_count = 0

    try:
        async with client.stream("POST", url, json=payload) as response:
            response.raise_for_status()
            event_name = None
            event_data = None

            async for line in response.aiter_lines():
                if line.startswith("event:"):
                    event_name = line.replace("event:", "", 1).strip()
                elif line.startswith("data:"):
                    event_data = line.replace("data:", "", 1).strip()
                elif line == "":
                    if not event_name or not event_data:
                        event_name = None
                        event_data = None
                        continue
                    payload = json.loads(event_data)
                    if event_name == "ttft" and ttft is None:
                        ttft = time.perf_counter() - start
                    if event_name == "token":
                        if ttft is None:
                            ttft = time.perf_counter() - start
                        token_count += max(1, len(payload.get("text", "").split()))
                    if event_name == "done":
                        break
                    event_name = None
                    event_data = None

        total = time.perf_counter() - start
        tokens_per_second = token_count / total if total > 0 else 0.0
        return {
            "id": request_id,
            "success": True,
            "ttft": ttft or total,
            "total": total,
            "tokens_per_second": tokens_per_second,
        }
    except Exception as exc:  # noqa: BLE001
        return {
            "id": request_id,
            "success": False,
            "error": str(exc),
        }


async def worker(
    name: int,
    client: httpx.AsyncClient,
    url: str,
    prompt: str,
    counter: asyncio.Lock,
    remaining: list[int],
    results: list[dict[str, Any]],
) -> None:
    while True:
        async with counter:
            if remaining[0] <= 0:
                return
            remaining[0] -= 1
            request_id = remaining[0]
        result = await run_request(client, url, prompt, request_id)
        results.append(result)


async def run_benchmark(args: argparse.Namespace) -> dict[str, Any]:
    prompt = "Explain what this system is and why vLLM matters."
    if args.prompt_file:
        prompt = Path(args.prompt_file).read_text().strip()

    timeout = httpx.Timeout(args.timeout)
    limits = httpx.Limits(max_keepalive_connections=args.concurrency)
    results: list[dict[str, Any]] = []
    remaining = [args.requests]
    lock = asyncio.Lock()

    async with httpx.AsyncClient(timeout=timeout, limits=limits) as client:
        tasks = [
            asyncio.create_task(worker(i, client, args.url, prompt, lock, remaining, results))
            for i in range(args.concurrency)
        ]
        await asyncio.gather(*tasks)

    success = [r for r in results if r.get("success")]
    failures = [r for r in results if not r.get("success")]
    ttft_values = [r["ttft"] for r in success]
    total_values = [r["total"] for r in success]
    tokens_per_second = [r["tokens_per_second"] for r in success]

    summary = {
        "requests": args.requests,
        "concurrency": args.concurrency,
        "success": len(success),
        "errors": len(failures),
        "ttft_p50_ms": round(percentile(ttft_values, 0.50) * 1000, 2),
        "ttft_p95_ms": round(percentile(ttft_values, 0.95) * 1000, 2),
        "latency_p50_ms": round(percentile(total_values, 0.50) * 1000, 2),
        "latency_p95_ms": round(percentile(total_values, 0.95) * 1000, 2),
        "avg_tokens_per_sec": round(
            statistics.mean(tokens_per_second) if tokens_per_second else 0.0, 2
        ),
    }

    if args.json_out:
        Path(args.json_out).write_text(json.dumps(summary, indent=2))

    if args.show_errors and failures:
        print("Sample errors:")
        for item in failures[: args.show_errors]:
            print(f"- {item.get('error')}")

    return summary


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Stream benchmark for /query/stream")
    parser.add_argument(
        "--url",
        default="http://localhost:8000/query/stream",
        help="Streaming endpoint URL",
    )
    parser.add_argument("--concurrency", type=int, default=10, help="Concurrent requests")
    parser.add_argument("--requests", type=int, default=100, help="Total requests")
    parser.add_argument("--prompt-file", help="Optional prompt file")
    parser.add_argument("--json-out", help="Write summary JSON to file")
    parser.add_argument("--timeout", type=int, default=120, help="Request timeout seconds")
    parser.add_argument(
        "--show-errors",
        type=int,
        default=0,
        help="Print up to N error messages",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    summary = asyncio.run(run_benchmark(args))
    print(json.dumps(summary, indent=2))


if __name__ == "__main__":
    main()
"""

def chunk_text(text: str, chunk_size: int = 800, overlap: int = 120) -> list[str]:
    if len(text) <= chunk_size:
        return [text]
    chunks = []
    start = 0
    while start < len(text):
        end = min(len(text), start + chunk_size)
        chunks.append(text[start:end])
        if end == len(text):
            break
        start = max(end - overlap, 0)
    return chunks


def extract_text_from_pdf(data: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(data))
        pages = []
        for page in reader.pages:
            pages.append(page.extract_text() or "")
        return "\n".join(pages)
    except Exception:  # noqa: BLE001
        return data.decode("utf-8", errors="ignore")


def extract_text_from_docx(data: bytes) -> str:
    try:
        doc = DocxDocument(io.BytesIO(data))
        return "\n".join(para.text for para in doc.paragraphs)
    except Exception:  # noqa: BLE001
        return data.decode("utf-8", errors="ignore")


def extract_text_from_html(data: bytes) -> str:
    try:
        soup = BeautifulSoup(data.decode("utf-8", errors="ignore"), "html.parser")
        return soup.get_text(separator=" ", strip=True)
    except Exception:  # noqa: BLE001
        return data.decode("utf-8", errors="ignore")


def load_text_from_upload(upload: Any) -> str:
    data = upload.file.read()
    if upload.filename and upload.filename.lower().endswith(".pdf"):
        return extract_text_from_pdf(data)
    if upload.filename and upload.filename.lower().endswith(".docx"):
        return extract_text_from_docx(data)
    if upload.filename and upload.filename.lower().endswith((".html", ".htm")):
        return extract_text_from_html(data)
    return data.decode("utf-8", errors="ignore")


def load_text_from_url(url: str, timeout: int = 10) -> str:
    response = requests.get(url, timeout=timeout)
    response.raise_for_status()
    content_type = response.headers.get("content-type", "")
    if "text/html" in content_type:
        soup = BeautifulSoup(response.text, "html.parser")
        return soup.get_text(separator=" ", strip=True)
    return response.text


class TimingTracker:
    def __init__(self, maxlen: int = 200) -> None:
        self.samples: dict[str, deque[float]] = defaultdict(lambda: deque(maxlen=maxlen))

    def record(self, name: str, value: float) -> None:
        self.samples[name].append(value)

    def summarize(self) -> dict[str, dict[str, float]]:
        summary: dict[str, dict[str, float]] = {}
        for key, values in self.samples.items():
            if not values:
                continue
            ordered = sorted(values)
            avg = sum(values) / len(values)
            p95 = ordered[int(len(ordered) * 0.95) - 1] if len(ordered) > 1 else ordered[0]
            summary[key] = {
                "avg_ms": round(avg * 1000, 2),
                "p95_ms": round(p95 * 1000, 2),
                "count": len(values),
            }
        return summary


@serve.deployment(ray_actor_options={"num_cpus": 1})
class RagApp:
    def __init__(self) -> None:
        configure_logging()
        self.logger = logging.getLogger("rag-app")
        self.request_counter = Counter(
            "rag_requests_total",
            "Total requests",
            ["endpoint"],
        )
        self.error_counter = Counter(
            "rag_errors_total",
            "Total errors",
            ["endpoint"],
        )
        self.latency_histogram = Histogram(
            "rag_latency_seconds",
            "Latency in seconds",
            ["stage"],
        )
        self.ttft_histogram = Histogram(
            "rag_ttft_seconds",
            "Time-to-first-token in seconds",
        )
        self.tokens_per_second_histogram = Histogram(
            "rag_tokens_per_second",
            "Tokens per second (estimated)",
        )
        self.token_counter = Counter(
            "rag_tokens_total",
            "Total tokens generated (estimated)",
        )
        self.timings = TimingTracker()
        self.use_embeddings = env_flag("RAG_USE_EMBEDDINGS", "true")
        self.qdrant_url = os.getenv("QDRANT_URL", "")
        self.qdrant_collection = os.getenv("QDRANT_COLLECTION", "rag-documents")
        self.embedding_model = os.getenv(
            "EMBEDDING_MODEL_ID",
            "sentence-transformers/all-MiniLM-L6-v2",
        )
        self.vllm_base_url = os.getenv("VLLM_BASE_URL", "http://vllm:8000")
        self.vllm_model = os.getenv("VLLM_MODEL", "Qwen/Qwen2.5-7B-Instruct")
        self.vllm_max_tokens = int(os.getenv("VLLM_MAX_TOKENS", "512"))
        self.vllm_temperature = float(os.getenv("VLLM_TEMPERATURE", "0.2"))
        self.vllm_top_p = float(os.getenv("VLLM_TOP_P", "0.95"))
        self.vllm_timeout_seconds = int(os.getenv("VLLM_TIMEOUT_SECONDS", "30"))
        self.max_history = int(os.getenv("RAG_MAX_HISTORY", "6"))
        self.top_k = int(os.getenv("RAG_TOP_K", "4"))
        self.sessions: dict[str, list[dict[str, str]]] = {}
        self.ingest_index: dict[str, set[str]] = defaultdict(set)
        self.provider = os.getenv("RAG_PROVIDER", "unknown")
        self.kube_namespace = os.getenv("KUBE_NAMESPACE") or os.getenv("KUBERNETES_NAMESPACE", "rag-app")
        self.bench_target_url = os.getenv(
            "BENCH_TARGET_URL",
            "http://rag-app-rag-app-backend:8000/query/stream",
        )
        self.bench_image = os.getenv("BENCH_IMAGE", "python:3.11-slim")
        self.bench_ttl_seconds = int(os.getenv("BENCH_TTL_SECONDS", "3600"))
        self._kube_api = self._build_kube_api()

        self.document_store = self._build_document_store()
        if self.qdrant_url:
            self.use_embeddings = True
        self.retriever = self._build_retriever()
        self.document_embedder = self._build_document_embedder()
        self.query_embedder = self._build_query_embedder()
        self.vllm = self._build_vllm_client()
        self._embedder_lock = threading.Lock()
        self._embedder_ready = {"document": False, "query": False}
        self._warm_up_embedders()

    def _warm_up_embedders(self) -> None:
        for name, embedder in (
            ("document", self.document_embedder),
            ("query", self.query_embedder),
        ):
            if embedder is None:
                continue
            warm_up = getattr(embedder, "warm_up", None)
            if callable(warm_up):
                try:
                    warm_up()
                    self._embedder_ready[name] = True
                except Exception as exc:  # noqa: BLE001
                    self.logger.warning(
                        "embedder_warmup_failed",
                        extra={"embedder": name, "error": str(exc)},
                    )
                    raise

    def _ensure_query_embedder_ready(self) -> None:
        if not self.query_embedder:
            return
        if self._embedder_ready.get("query"):
            return
        with self._embedder_lock:
            if self._embedder_ready.get("query"):
                return
            warm_up = getattr(self.query_embedder, "warm_up", None)
            if callable(warm_up):
                try:
                    warm_up()
                except Exception:  # noqa: BLE001
                    raise
            self._embedder_ready["query"] = True

    def _ensure_document_embedder_ready(self) -> None:
        if not self.document_embedder:
            return
        if self._embedder_ready.get("document"):
            return
        with self._embedder_lock:
            if self._embedder_ready.get("document"):
                return
            warm_up = getattr(self.document_embedder, "warm_up", None)
            if callable(warm_up):
                try:
                    warm_up()
                except Exception:  # noqa: BLE001
                    raise
            self._embedder_ready["document"] = True

    def _build_document_store(self) -> Any:
        if self.qdrant_url:
            return QdrantDocumentStore(
                url=self.qdrant_url,
                index=self.qdrant_collection,
            )
        return InMemoryDocumentStore()

    def _build_retriever(self) -> Any:
        if self.use_embeddings:
            if self.qdrant_url:
                return QdrantEmbeddingRetriever(document_store=self.document_store)
            return InMemoryEmbeddingRetriever(document_store=self.document_store)
        if isinstance(self.document_store, InMemoryDocumentStore):
            return InMemoryBM25Retriever(document_store=self.document_store)
        raise ValueError("BM25 retriever is only supported with in-memory store.")

    def _build_document_embedder(self) -> Any | None:
        if not self.use_embeddings:
            return None
        return SentenceTransformersDocumentEmbedder(model=self.embedding_model)

    def _build_query_embedder(self) -> Any | None:
        if not self.use_embeddings:
            return None
        return SentenceTransformersTextEmbedder(model=self.embedding_model)

    def _build_vllm_client(self) -> VllmStreamingGenerator:
        return VllmStreamingGenerator(
            base_url=self.vllm_base_url,
            model=self.vllm_model,
            max_tokens=self.vllm_max_tokens,
            temperature=self.vllm_temperature,
            top_p=self.vllm_top_p,
            timeout_seconds=self.vllm_timeout_seconds,
        )

    def _build_kube_api(self) -> dict[str, Any] | None:
        host = os.getenv("KUBERNETES_SERVICE_HOST")
        port = os.getenv("KUBERNETES_SERVICE_PORT")
        if not host or not port:
            return None
        token_path = "/var/run/secrets/kubernetes.io/serviceaccount/token"
        ca_path = "/var/run/secrets/kubernetes.io/serviceaccount/ca.crt"
        if not os.path.exists(token_path):
            return None
        token = Path(token_path).read_text().strip()
        verify = ca_path if os.path.exists(ca_path) else True
        return {
            "base_url": f"https://{host}:{port}",
            "token": token,
            "verify": verify,
        }

    def _kube_request(
        self,
        method: str,
        path: str,
        json_body: dict[str, Any] | None = None,
        headers: dict[str, str] | None = None,
    ) -> requests.Response:
        if not self._kube_api:
            raise ValueError("Kubernetes API is not available in this environment.")
        url = f"{self._kube_api['base_url']}{path}"
        request_headers = {
            "Authorization": f"Bearer {self._kube_api['token']}",
            "Accept": "application/json",
        }
        if headers:
            request_headers.update(headers)
        response = requests.request(
            method,
            url,
            headers=request_headers,
            json=json_body,
            timeout=15,
            verify=self._kube_api["verify"],
        )
        if not response.ok:
            raise ValueError(f"kubernetes_api_error: {response.status_code} {response.text}")
        return response

    def _coerce_int(self, value: Any, default: int, minimum: int = 1) -> int:
        try:
            parsed = int(value)
        except (TypeError, ValueError):
            return default
        return parsed if parsed >= minimum else default

    def _benchmark_configmap(self, name: str) -> dict[str, Any]:
        return {
            "apiVersion": "v1",
            "kind": "ConfigMap",
            "metadata": {
                "name": name,
                "namespace": self.kube_namespace,
                "labels": {"app": "rag-stream-bench"},
            },
            "data": {
                "requirements.txt": BENCH_REQUIREMENTS,
                "stream_bench.py": BENCH_SCRIPT,
            },
        }

    def _benchmark_job(
        self,
        name: str,
        configmap_name: str,
        concurrency: int,
        requests_total: int,
        timeout: int,
        show_errors: int,
    ) -> dict[str, Any]:
        command = (
            "python -m pip install --no-cache-dir -r /bench/requirements.txt && "
            "python /bench/stream_bench.py "
            "--url \"$BENCH_URL\" "
            "--concurrency \"$BENCH_CONCURRENCY\" "
            "--requests \"$BENCH_REQUESTS\" "
            "--timeout \"$BENCH_TIMEOUT\" "
            "--show-errors \"$BENCH_SHOW_ERRORS\""
        )
        return {
            "apiVersion": "batch/v1",
            "kind": "Job",
            "metadata": {
                "name": name,
                "namespace": self.kube_namespace,
                "labels": {"app": "rag-stream-bench"},
            },
            "spec": {
                "backoffLimit": 0,
                "ttlSecondsAfterFinished": self.bench_ttl_seconds,
                "template": {
                    "metadata": {"labels": {"app": "rag-stream-bench"}},
                    "spec": {
                        "restartPolicy": "Never",
                        "containers": [
                            {
                                "name": "bench",
                                "image": self.bench_image,
                                "imagePullPolicy": "IfNotPresent",
                                "env": [
                                    {"name": "PIP_DISABLE_PIP_VERSION_CHECK", "value": "1"},
                                    {"name": "PYTHONUNBUFFERED", "value": "1"},
                                    {"name": "BENCH_URL", "value": self.bench_target_url},
                                    {
                                        "name": "BENCH_CONCURRENCY",
                                        "value": str(concurrency),
                                    },
                                    {"name": "BENCH_REQUESTS", "value": str(requests_total)},
                                    {"name": "BENCH_TIMEOUT", "value": str(timeout)},
                                    {"name": "BENCH_SHOW_ERRORS", "value": str(show_errors)},
                                ],
                                "command": ["/bin/sh", "-c", command],
                                "volumeMounts": [{"name": "bench", "mountPath": "/bench"}],
                            }
                        ],
                        "volumes": [
                            {
                                "name": "bench",
                                "configMap": {"name": configmap_name},
                            }
                        ],
                    },
                },
            },
        }

    def _benchmark_run(self, payload: dict[str, Any]) -> dict[str, Any]:
        concurrency = self._coerce_int(payload.get("concurrency"), 10)
        requests_total = self._coerce_int(payload.get("requests"), 100)
        timeout = self._coerce_int(payload.get("timeout"), 120)
        show_errors = self._coerce_int(payload.get("show_errors"), 3, minimum=0)
        run_id = uuid4().hex[:8]
        job_name = f"rag-stream-bench-{run_id}"
        configmap_name = f"rag-stream-bench-{run_id}"

        configmap_body = self._benchmark_configmap(configmap_name)
        job_body = self._benchmark_job(
            job_name,
            configmap_name,
            concurrency,
            requests_total,
            timeout,
            show_errors,
        )

        self._kube_request(
            "POST",
            f"/api/v1/namespaces/{self.kube_namespace}/configmaps",
            json_body=configmap_body,
        )
        job_response = self._kube_request(
            "POST",
            f"/apis/batch/v1/namespaces/{self.kube_namespace}/jobs",
            json_body=job_body,
        ).json()

        job_uid = job_response.get("metadata", {}).get("uid")
        if job_uid:
            patch_body = {
                "metadata": {
                    "ownerReferences": [
                        {
                            "apiVersion": "batch/v1",
                            "kind": "Job",
                            "name": job_name,
                            "uid": job_uid,
                        }
                    ]
                }
            }
            self._kube_request(
                "PATCH",
                f"/api/v1/namespaces/{self.kube_namespace}/configmaps/{configmap_name}",
                json_body=patch_body,
                headers={"Content-Type": "application/merge-patch+json"},
            )

        return {
            "job_name": job_name,
            "namespace": self.kube_namespace,
            "concurrency": concurrency,
            "requests": requests_total,
            "timeout": timeout,
            "show_errors": show_errors,
        }

    def _benchmark_status(self, job_name: str) -> dict[str, Any]:
        job_response = self._kube_request(
            "GET",
            f"/apis/batch/v1/namespaces/{self.kube_namespace}/jobs/{job_name}",
        ).json()
        status = job_response.get("status", {})
        phase = "pending"
        if status.get("active"):
            phase = "running"
        if status.get("succeeded"):
            phase = "succeeded"
        if status.get("failed"):
            phase = "failed"
        return {
            "job_name": job_name,
            "phase": phase,
            "start_time": status.get("startTime"),
            "completion_time": status.get("completionTime"),
            "failed": status.get("failed", 0),
        }

    def _benchmark_logs(self, job_name: str) -> str:
        label_selector = quote(f"job-name={job_name}")
        pods_response = self._kube_request(
            "GET",
            f"/api/v1/namespaces/{self.kube_namespace}/pods?labelSelector={label_selector}",
            headers={"Accept": "application/json"},
        )
        pods = pods_response.json().get("items", [])
        if not pods:
            raise ValueError("benchmark_logs_not_ready")
        pod_name = pods[0]["metadata"]["name"]
        logs_response = self._kube_request(
            "GET",
            f"/api/v1/namespaces/{self.kube_namespace}/pods/{pod_name}/log",
            headers={"Accept": "text/plain"},
        )
        return logs_response.text

    def _make_document(self, content: str, meta: dict[str, Any], key: str | None = None) -> Document:
        doc_meta = dict(meta)
        if key:
            doc_meta.setdefault("ingest_key", key)
        return Document(id=str(uuid4()), content=content, meta=doc_meta)

    def _track_documents(self, key: str | None, docs: list[Document]) -> None:
        if not key:
            return
        for doc in docs:
            if doc.id:
                self.ingest_index[key].add(doc.id)

    def _get_session_history(self, session_id: str | None) -> tuple[str, list[dict[str, str]]]:
        if not session_id:
            session_id = str(uuid4())
        history = self.sessions.get(session_id, [])
        return session_id, history

    def _update_session(self, session_id: str, role: str, content: str) -> None:
        history = self.sessions.setdefault(session_id, [])
        history.append({"role": role, "content": content})
        if len(history) > self.max_history:
            self.sessions[session_id] = history[-self.max_history :]

    async def healthz(self) -> dict[str, str]:
        self.request_counter.labels("healthz").inc()
        return {"status": "ok"}

    async def metrics(self) -> PlainTextResponse:
        self.request_counter.labels("metrics").inc()
        return PlainTextResponse(generate_latest(), media_type=CONTENT_TYPE_LATEST)

    async def stats(self) -> dict[str, Any]:
        self.request_counter.labels("stats").inc()
        return {
            "provider": self.provider,
            "sessions": len(self.sessions),
            "timings": self.timings.summarize(),
        }

    async def ingest(
        self,
        files: list[Any] | None,
        payload: dict[str, Any] | None,
    ) -> dict[str, Any]:
        self.request_counter.labels("ingest").inc()
        payload = payload or {}
        documents: list[Document] = []
        errors: list[str] = []

        start_time = time.perf_counter()

        if files:
            for upload in files:
                try:
                    content = load_text_from_upload(upload)
                    file_key = upload.filename or f"upload-{uuid4()}"
                    file_docs: list[Document] = []
                    for chunk in chunk_text(content):
                        file_docs.append(
                            self._make_document(
                                chunk,
                                {"filename": upload.filename, "source": "file"},
                                key=file_key,
                            )
                        )
                    documents.extend(file_docs)
                    self._track_documents(file_key, file_docs)
                except Exception as exc:  # noqa: BLE001
                    errors.append(str(exc))

        for item in payload.get("documents", []):
            meta = dict(item.get("meta", {}))
            doc_key = meta.get("filename") or meta.get("ingest_key")
            doc = self._make_document(item.get("content", ""), meta, key=doc_key)
            documents.append(doc)
            if doc_key:
                self._track_documents(doc_key, [doc])

        for index, text in enumerate(payload.get("texts", [])):
            text_key = f"text:{index}"
            text_docs: list[Document] = []
            for chunk in chunk_text(text):
                text_docs.append(
                    self._make_document(chunk, {"source": "text"}, key=text_key)
                )
            documents.extend(text_docs)
            self._track_documents(text_key, text_docs)

        for url in payload.get("urls", []):
            try:
                content = load_text_from_url(url)
                url_docs: list[Document] = []
                for chunk in chunk_text(content):
                    url_docs.append(
                        self._make_document(
                            chunk,
                            {"source": "url", "url": url},
                            key=url,
                        )
                    )
                documents.extend(url_docs)
                self._track_documents(url, url_docs)
            except Exception as exc:  # noqa: BLE001
                errors.append(f"{url}: {exc}")

        sitemap_url = payload.get("sitemap_url")
        if sitemap_url:
            try:
                response = requests.get(sitemap_url, timeout=10)
                response.raise_for_status()
                soup = BeautifulSoup(response.text, "xml")
                for loc in soup.find_all("loc"):
                    url = loc.text.strip()
                    if not url:
                        continue
                    try:
                        content = load_text_from_url(url)
                        sitemap_docs: list[Document] = []
                        for chunk in chunk_text(content):
                            sitemap_docs.append(
                                self._make_document(
                                    chunk,
                                    {"source": "sitemap", "url": url},
                                    key=f"sitemap:{url}",
                                )
                            )
                        documents.extend(sitemap_docs)
                        self._track_documents(f"sitemap:{url}", sitemap_docs)
                    except Exception as exc:  # noqa: BLE001
                        errors.append(f"{url}: {exc}")
            except Exception as exc:  # noqa: BLE001
                errors.append(f"sitemap: {exc}")

        if not documents:
            return {"ingested": 0, "errors": errors}

        if self.use_embeddings and self.document_embedder:
            self._ensure_document_embedder_ready()
            documents = self.document_embedder.run(documents=documents)["documents"]

        self.document_store.write_documents(documents)
        duration = time.perf_counter() - start_time
        self.latency_histogram.labels("ingest").observe(duration)
        self.timings.record("ingest", duration)
        self.logger.info(
            "ingested_documents",
            extra={"count": len(documents), "errors": len(errors)},
        )
        return {"ingested": len(documents), "errors": errors}

    async def delete(self, payload: dict[str, Any]) -> dict[str, Any]:
        self.request_counter.labels("delete").inc()
        if not hasattr(self.document_store, "delete_documents"):
            return {"deleted": 0, "error": "delete is not supported by document store"}

        delete_all = bool(payload.get("all"))
        filenames = payload.get("filenames", [])
        keys = payload.get("keys", [])
        document_ids = payload.get("document_ids", [])
        ids: set[str] = set(document_ids)

        for name in filenames:
            ids.update(self.ingest_index.get(name, set()))
        for key in keys:
            ids.update(self.ingest_index.get(key, set()))

        try:
            if delete_all:
                self.document_store.delete_documents()
                self.ingest_index.clear()
                return {"deleted": "all"}
            if not ids:
                return {"deleted": 0, "error": "no matching documents"}
            self.document_store.delete_documents(list(ids))
            for name in filenames:
                self.ingest_index.pop(name, None)
            for key in keys:
                self.ingest_index.pop(key, None)
            return {"deleted": len(ids)}
        except Exception as exc:  # noqa: BLE001
            self.error_counter.labels("delete").inc()
            return {"deleted": 0, "error": str(exc)}

    async def list_documents(self) -> dict[str, Any]:
        self.request_counter.labels("documents").inc()
        items = [
            {"key": key, "count": len(ids)} for key, ids in sorted(self.ingest_index.items())
        ]
        return {"items": items}

    async def query(self, payload: dict[str, Any]) -> dict[str, Any]:
        self.request_counter.labels("query").inc()
        query = payload.get("query", "")
        if not query:
            return {"answers": [], "documents": []}

        session_id, history = self._get_session_history(payload.get("session_id"))
        if payload.get("history"):
            history = payload["history"]

        self._update_session(session_id, "user", query)

        retrieval_start = time.perf_counter()
        if self.use_embeddings and self.query_embedder:
            self._ensure_query_embedder_ready()
            embedding = self.query_embedder.run(text=query)["embedding"]
            result = self.retriever.run(query_embedding=embedding, top_k=self.top_k)
        else:
            result = self.retriever.run(query=query, top_k=self.top_k)
        documents = result.get("documents", [])
        retrieval_time = time.perf_counter() - retrieval_start

        prompt = self._build_prompt(query, history, documents)

        generation_start = time.perf_counter()
        ttft_start = time.perf_counter()
        answer = ""
        try:
            answer = await self.vllm.complete_chat(prompt)
        except Exception as exc:  # noqa: BLE001
            self.error_counter.labels("query").inc()
            answer = f"Generation failed: {exc}"
        generation_time = time.perf_counter() - generation_start
        ttft = time.perf_counter() - ttft_start
        # Token estimation uses whitespace splits to avoid tokenizer overhead.
        token_count = max(1, len(answer.split())) if answer else 0
        tokens_per_second = token_count / generation_time if generation_time > 0 else 0.0

        self._update_session(session_id, "assistant", answer)
        self.timings.record("retrieval", retrieval_time)
        self.timings.record("generation", generation_time)
        self.timings.record("ttft", ttft)
        self.latency_histogram.labels("retrieval").observe(retrieval_time)
        self.latency_histogram.labels("generation").observe(generation_time)
        self.ttft_histogram.observe(ttft)
        if token_count:
            self.token_counter.inc(token_count)
            self.tokens_per_second_histogram.observe(tokens_per_second)

        total_time = retrieval_time + generation_time
        self.latency_histogram.labels("total").observe(total_time)

        self.logger.info(
            "query",
            extra={
                "query": query,
                "documents": len(documents),
                "session_id": session_id,
                "retrieval_ms": round(retrieval_time * 1000, 2),
                "generation_ms": round(generation_time * 1000, 2),
            },
        )
        return {
            "session_id": session_id,
            "answers": [{"answer": answer}],
            "documents": [
                {
                    "content": doc.content,
                    "meta": doc.meta,
                    "score": doc.score,
                }
                for doc in documents
            ],
            "timings": {
                "retrieval_ms": round(retrieval_time * 1000, 2),
                "generation_ms": round(generation_time * 1000, 2),
                "total_ms": round(total_time * 1000, 2),
                "ttft_ms": round(ttft * 1000, 2),
                "tokens_per_second": round(tokens_per_second, 2),
                "tokens_estimated": token_count,
            },
            "history": self.sessions.get(session_id, []),
        }

    async def query_stream(self, payload: dict[str, Any]) -> StreamingResponse:
        self.request_counter.labels("query_stream").inc()
        query = payload.get("query", "")
        if not query:
            return StreamingResponse(
                self._stream_events(
                    [
                        {
                            "event": "error",
                            "data": {"message": "query is required"},
                        }
                    ]
                ),
                media_type="text/event-stream",
            )

        request_id = uuid4().hex
        session_id, history = self._get_session_history(payload.get("session_id"))
        if payload.get("history"):
            history = payload["history"]
        replica_id = os.getenv("HOSTNAME", "unknown")
        model_id = (
            os.getenv("VLLM_MODEL_ID")
            or os.getenv("MODEL_ID")
            or os.getenv("VLLM_MODEL")
            or "unknown"
        )

        self._update_session(session_id, "user", query)

        retrieval_start = time.perf_counter()
        if self.use_embeddings and self.query_embedder:
            self._ensure_query_embedder_ready()
            embedding = self.query_embedder.run(text=query)["embedding"]
            result = self.retriever.run(query_embedding=embedding, top_k=self.top_k)
        else:
            result = self.retriever.run(query=query, top_k=self.top_k)
        documents = result.get("documents", [])
        retrieval_time = time.perf_counter() - retrieval_start
        k = len(documents)

        prompt = self._build_prompt(query, history, documents)

        async def event_stream() -> AsyncIterator[bytes]:
            # SSE event contract:
            # meta -> retrieval docs + timings, ttft -> time to first token,
            # token -> incremental delta, done -> final timings + citations.
            yield sse(
                "meta",
                {
                    "session_id": session_id,
                    "request_id": request_id,
                    "replica_id": replica_id,
                    "model_id": model_id,
                    "k": k,
                    "documents": [
                        {
                            "content": doc.content,
                            "meta": doc.meta,
                            "score": doc.score,
                        }
                        for doc in documents
                    ],
                    "timings": {
                        "retrieval_ms": round(retrieval_time * 1000, 2),
                    },
                },
            )

            generation_start = time.perf_counter()
            server_start = generation_start
            token_count = 0
            server_first_token_at: float | None = None

            try:
                async for delta in self.vllm.stream_chat(prompt):
                    if server_first_token_at is None:
                        server_first_token_at = time.perf_counter()
                        ttft_value = server_first_token_at - server_start
                        self.ttft_histogram.observe(ttft_value)
                        self.timings.record("ttft", ttft_value)
                        yield sse(
                            "ttft",
                            {
                                "ttft_ms": round(ttft_value * 1000, 2),
                                "request_id": request_id,
                                "session_id": session_id,
                            },
                        )

                    token_count += 1
                    yield sse("token", {"text": delta})

                generation_time = time.perf_counter() - generation_start
                total_ms = (time.perf_counter() - server_start) * 1000
                ttft_ms = (
                    (server_first_token_at - server_start) * 1000
                    if server_first_token_at is not None
                    else None
                )
                stream_duration_sec = (
                    (time.perf_counter() - server_first_token_at)
                    if server_first_token_at is not None
                    else None
                )
                tokens_per_second = (
                    token_count / stream_duration_sec
                    if stream_duration_sec and stream_duration_sec > 0
                    else None
                )

                self.token_counter.inc(token_count)
                if tokens_per_second is not None:
                    self.tokens_per_second_histogram.observe(tokens_per_second)
                self.timings.record("retrieval", retrieval_time)
                self.timings.record("generation", generation_time)
                self.latency_histogram.labels("retrieval").observe(retrieval_time)
                self.latency_histogram.labels("generation").observe(generation_time)
                total_time = retrieval_time + generation_time
                self.latency_histogram.labels("total").observe(total_time)

                yield sse(
                    "done",
                    {
                        "session_id": session_id,
                        "request_id": request_id,
                        "replica_id": replica_id,
                        "model_id": model_id,
                        "k": k,
                        "documents": [
                            {
                                "content": doc.content,
                                "meta": doc.meta,
                                "score": doc.score,
                            }
                            for doc in documents
                        ],
                        "timings": {
                            "ttft_ms": round(ttft_ms, 2) if ttft_ms is not None else None,
                            "total_ms": round(total_ms, 2),
                        },
                        "token_count": token_count,
                        "tokens_per_sec": round(tokens_per_second, 2)
                        if tokens_per_second is not None
                        else None,
                    },
                )
            except Exception as exc:  # noqa: BLE001
                self.error_counter.labels("query_stream").inc()
                yield sse(
                    "error",
                    {
                        "message": "Streaming failed",
                        "session_id": session_id,
                        "request_id": request_id,
                    },
                )

        return StreamingResponse(event_stream(), media_type="text/event-stream")

    def _build_prompt(
        self,
        query: str,
        history: list[dict[str, str]],
        documents: list[Document],
    ) -> str:
        prompt_context = "\n\n".join(
            f"[{index + 1}] {doc.content}" for index, doc in enumerate(documents)
        )
        history_text = "\n".join(
            f"{message['role']}: {message['content']}"
            for message in history[-self.max_history :]
        )
        return (
            "You are a helpful RAG assistant. Use the context to answer.\n\n"
            f"Context:\n{prompt_context}\n\n"
            f"Conversation:\n{history_text}\n\n"
            f"Question: {query}\nAnswer:"
        )

    def _sse_event(self, name: str, data: dict[str, Any]) -> bytes:
        return f"event: {name}\ndata: {json.dumps(data)}\n\n".encode()

    def _stream_events(self, events: list[dict[str, Any]]) -> Any:
        for event in events:
            payload = event.get("data", {})
            event_name = event.get("event", "message")
            yield f"event: {event_name}\n".encode()
            yield f"data: {json.dumps(payload)}\n\n".encode()

    async def __call__(self, request: Request) -> Response:
        path = request.url.path
        method = request.method.upper()

        if path == "/healthz" and method == "GET":
            return JSONResponse(await self.healthz())

        if path == "/metrics" and method == "GET":
            return await self.metrics()

        if path == "/stats" and method == "GET":
            return JSONResponse(await self.stats())

        if path == "/ingest" and method == "POST":
            content_type = request.headers.get("content-type", "")
            files = None
            payload = None
            if "multipart/form-data" in content_type:
                form = await request.form()
                files = form.getlist("files")
            else:
                try:
                    payload = await request.json()
                except ValueError:
                    payload = None
            return JSONResponse(await self.ingest(files=files, payload=payload))

        if path == "/delete" and method == "POST":
            payload = await request.json()
            return JSONResponse(await self.delete(payload))

        if path == "/documents" and method == "GET":
            return JSONResponse(await self.list_documents())

        if path == "/benchmark/run" and method == "POST":
            try:
                payload = await request.json()
            except ValueError:
                payload = {}
            try:
                result = self._benchmark_run(payload or {})
                return JSONResponse(result)
            except ValueError as exc:
                return JSONResponse({"error": str(exc)}, status_code=400)
            except Exception as exc:  # noqa: BLE001
                self.logger.error("benchmark_run_failed", extra={"error": str(exc)})
                return JSONResponse({"error": "benchmark_run_failed"}, status_code=500)

        if path == "/benchmark/status" and method == "GET":
            job_name = request.query_params.get("job")
            if not job_name:
                return JSONResponse({"error": "job is required"}, status_code=400)
            try:
                result = self._benchmark_status(job_name)
                return JSONResponse(result)
            except ValueError as exc:
                return JSONResponse({"error": str(exc)}, status_code=400)
            except Exception as exc:  # noqa: BLE001
                self.logger.error("benchmark_status_failed", extra={"error": str(exc)})
                return JSONResponse({"error": "benchmark_status_failed"}, status_code=500)

        if path == "/benchmark/logs" and method == "GET":
            job_name = request.query_params.get("job")
            if not job_name:
                return JSONResponse({"error": "job is required"}, status_code=400)
            try:
                logs = self._benchmark_logs(job_name)
                return PlainTextResponse(logs)
            except ValueError as exc:
                return JSONResponse({"error": str(exc)}, status_code=400)
            except Exception as exc:  # noqa: BLE001
                self.logger.error("benchmark_logs_failed", extra={"error": str(exc)})
                return JSONResponse({"error": "benchmark_logs_failed"}, status_code=500)

        if path == "/query" and method == "POST":
            payload = await request.json()
            return JSONResponse(await self.query(payload))

        if path == "/query/stream" and method == "POST":
            payload = await request.json()
            return await self.query_stream(payload)

        return JSONResponse({"error": "not_found"}, status_code=404)


deployment = RagApp.bind()
